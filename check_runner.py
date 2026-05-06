"""
check_runner.py — Автоматическая проверка журналов для GitHub Actions
Каждый день 21:00 Актау (16:00 UTC)
Отправляет краткое сообщение + Word файл в Telegram
"""
import io
import json
import os
import sys
import time
from datetime import datetime, timedelta

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ── Конфигурация из GitHub Secrets ────────────────────────────────────────────
TELEGRAM_TOKEN   = os.environ.get("TELEGRAM_TOKEN", "")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID", "")
TEACHERS_JSON    = os.environ.get("TEACHERS_JSON", "[]")

BASE_URL = "https://mektep.edu.kz"


# ── Telegram утилиты ──────────────────────────────────────────────────────────
def tg_send(text: str):
    if not TELEGRAM_TOKEN or not TELEGRAM_CHAT_ID:
        print("[TG]", text)
        return
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage"
    chunks = [text[i:i+4000] for i in range(0, len(text), 4000)]
    for chunk in chunks:
        try:
            requests.post(url, json={
                "chat_id": TELEGRAM_CHAT_ID,
                "text": chunk,
                "parse_mode": "HTML",
            }, timeout=15)
        except Exception as e:
            print(f"[TG] Ошибка: {e}")
        time.sleep(0.3)


def tg_send_document(file_bytes: bytes, filename: str, caption: str = ""):
    """Отправляет файл в Telegram."""
    if not TELEGRAM_TOKEN or not TELEGRAM_CHAT_ID:
        print(f"[TG] Файл сохранён локально: {filename}")
        with open(filename, "wb") as f:
            f.write(file_bytes)
        return
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument"
    try:
        resp = requests.post(url, data={
            "chat_id": TELEGRAM_CHAT_ID,
            "caption": caption,
            "parse_mode": "HTML",
        }, files={
            "document": (filename, file_bytes,
                         "application/vnd.openxmlformats-officedocument"
                         ".wordprocessingml.document")
        }, timeout=60)
        if resp.ok:
            print(f"[TG] Файл жіберілді: {filename}")
        else:
            print(f"[TG] Файл жіберу қатесі: {resp.text}")
    except Exception as e:
        print(f"[TG] Файл жіберу exception: {e}")


# ── Word helpers ──────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell(cells, i, text, bg=None, bold=False, center=True, size=8):
    cells[i].text = ""
    p = cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if bg:
        _set_cell_bg(cells[i], bg)


# ── Логика проверки ───────────────────────────────────────────────────────────
def make_session():
    s = requests.Session()
    s.headers.update({
        "User-Agent":      "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                           "AppleWebKit/537.36 (KHTML, like Gecko) "
                           "Chrome/120.0.0.0 Safari/537.36",
        "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "ru,kk;q=0.9,en;q=0.8",
        "Accept-Encoding": "gzip, deflate, br",
        "Connection":      "keep-alive",
    })
    adapter = requests.adapters.HTTPAdapter(
        max_retries=requests.adapters.Retry(total=3, backoff_factor=1)
    )
    s.mount("https://", adapter)
    return s


def login(session, login_val: str, password: str) -> bool:
    """
    Авторизация на mektep.edu.kz
    Форма: POST /ru/auth  поля: login, password, lang=ru
    """
    try:
        # Шаг 1: GET /ru/auth — получаем CSRF токен если есть
        r    = session.get(BASE_URL + "/ru/auth", timeout=15)
        soup = BeautifulSoup(r.text, "html.parser")
        tok  = soup.find("input", {"name": "_token"})

        # Шаг 2: POST авторизация с обязательным полем lang=ru
        r2 = session.post(
            BASE_URL + "/ru/auth",
            data={
                "login":    login_val,
                "password": password,
                "lang":     "ru",        # скрытое поле — обязательно!
                "_token":   tok["value"] if tok else "",
            },
            timeout=15,
            allow_redirects=True,
            headers={
                "Referer":      BASE_URL + "/ru/auth",
                "Origin":       BASE_URL,
                "Content-Type": "application/x-www-form-urlencoded",
            }
        )

        # Шаг 3: закрепляем сессию
        try:
            session.get(BASE_URL + "/ru", timeout=10, allow_redirects=True)
        except Exception:
            pass

        # Проверяем — если всё ещё на /auth значит пароль неверный
        final_url = r2.url.rstrip("/")
        if final_url.endswith("/ru/auth") or final_url.endswith("/kk/auth"):
            print(f"    [login] Авторизация сәтсіз — логин/пароль қате?")
            return False

        print(f"    [login] OK → {final_url}")
        return True

    except Exception as e:
        print(f"    [login] Қате: {e}")
        return False


def get_subjects(session) -> list:
    try:
        r    = session.get(BASE_URL + "/ru/grades", timeout=15,
                           allow_redirects=True)
        # Если нас перекинули на /auth — сессия не сохранилась
        if "/auth" in r.url:
            print(f"    [subjects] Сессия жоқ — /auth-қа өтті")
            return []
        soup = BeautifulSoup(r.text, "html.parser")
        out  = []
        for a in soup.select("a[href*='/ru/grades/']"):
            href = a.get("href", "")
            txt  = " ".join(a.get_text().split())
            if href not in ("/ru/grades", "") and txt:
                out.append({
                    "name": txt,
                    "url":  BASE_URL + href if href.startswith("/") else href
                })
        print(f"    [subjects] Табылды: {len(out)} пән")
        return out
    except Exception as e:
        print(f"    [subjects] Қате: {e}")
        return []


def check_subject(session, url: str, days_back: int = 14) -> list:
    """Возвращает список незаполненных дат."""
    today    = datetime.now()
    unfilled = []
    try:
        r    = session.get(url, timeout=20)
        soup = BeautifulSoup(r.text, "html.parser")
        table = soup.find("table")
        if not table:
            return []
        headers = table.find_all("th")
        date_cols = {}
        for idx, th in enumerate(headers):
            txt = th.get_text(strip=True)
            for fmt in ["%d.%m", "%d.%m.%Y"]:
                try:
                    dt = datetime.strptime(
                        txt + (f".{today.year}" if len(txt) <= 5 else ""),
                        "%d.%m.%Y"
                    )
                    if 0 <= (today - dt).days <= days_back:
                        date_cols[idx] = txt
                    break
                except ValueError:
                    continue
        rows = table.find_all("tr")[1:]
        for col_idx, date_str in date_cols.items():
            col_cells = []
            for row in rows:
                cells = row.find_all(["td", "th"])
                if col_idx < len(cells):
                    col_cells.append(cells[col_idx].get_text(strip=True))
            if col_cells and all(c in ("", "-", "н") for c in col_cells):
                unfilled.append(date_str)
    except Exception as e:
        print(f"    [check_subject] {e}")
    return unfilled


def check_teacher(teacher: dict) -> dict:
    name     = teacher.get("name", "—")
    login_v  = str(teacher.get("login", ""))
    password = str(teacher.get("password", ""))
    print(f"  👤 {name}")

    result = {"name": name, "login": login_v,
              "ok": False, "error": None, "unfilled": []}

    session = make_session()
    if not login(session, login_v, password):
        result["error"] = "Кіру мүмкін болмады"
        print("    ❌ Кіру қатесі")
        return result

    subjects = get_subjects(session)
    if not subjects:
        result["error"] = "Пәндер табылмады"
        return result

    all_ok = True
    for subj in subjects:
        dates = check_subject(session, subj["url"])
        if dates:
            all_ok = False
            result["unfilled"].append({"subject": subj["name"], "dates": dates})
            print(f"    ✘ {subj['name']}: {', '.join(dates)}")
        else:
            print(f"    ✓ {subj['name']}")
        time.sleep(0.5)

    result["ok"] = all_ok
    return result


# ── Генерация Word отчёта ─────────────────────────────────────────────────────
def build_word_report(results: list, check_time: datetime) -> bytes:
    doc = Document()

    # Поля
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(1.5)
    sec.left_margin = sec.right_margin = Cm(1.8)

    # Стиль
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(9)

    # ── Заголовок ──────────────────────────────────────────────────────────────
    t = doc.add_heading("", level=1)
    run = t.add_run("Электрондық журнал толтыру туралы есеп")
    run.font.size = Pt(13); run.bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph(
        f"Тексерілген күн: {check_time.strftime('%d.%m.%Y')}  |  "
        f"Уақыт: {check_time.strftime('%H:%M')} (Ақтау)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)

    # ── Статистика ─────────────────────────────────────────────────────────────
    total    = len(results)
    ok_cnt   = sum(1 for r in results if r["ok"] and not r["error"])
    bad_cnt  = sum(1 for r in results if r["unfilled"])
    err_cnt  = sum(1 for r in results if r["error"])

    stat_tbl = doc.add_table(rows=1, cols=4)
    stat_tbl.style = "Table Grid"
    stat_cells = stat_tbl.rows[0].cells
    stat_data = [
        ("Барлығы",          str(total),    "4472C4", "FFFFFF"),
        ("✅ Толтырылған",   str(ok_cnt),   "1B7A3E", "FFFFFF"),
        ("❌ Толтырылмаған", str(bad_cnt),  "C0392B", "FFFFFF"),
        ("⚠️ Қате",          str(err_cnt),  "E67E22", "FFFFFF"),
    ]
    for i, (label, val, bg, fg) in enumerate(stat_data):
        _set_cell_bg(stat_cells[i], bg)
        p2 = stat_cells[i].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r1 = p2.add_run(label + "\n")
        r1.font.size = Pt(8); r1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r2 = p2.add_run(val)
        r2.font.size = Pt(16); r2.bold = True
        r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Таблица проблемных ─────────────────────────────────────────────────────
    problems = [r for r in results if r["unfilled"] or r["error"]]

    if not problems:
        ok_p = doc.add_paragraph()
        ok_p.add_run("🎉 Барлық мұғалімдер журналды дұрыс толтырған!").bold = True
        ok_p.runs[0].font.color.rgb = RGBColor(0x1B,0x7A,0x3E)
        ok_p.runs[0].font.size = Pt(11)
        ok_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        hdr_p = doc.add_paragraph()
        hdr_p.add_run(f"Толтырылмаған немесе қатесі бар мұғалімдер — {len(problems)} адам:").bold = True
        hdr_p.paragraph_format.space_after = Pt(3)

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, (txt, bg) in enumerate([
            ("Мұғалімнің аты-жөні","1F3864"),
            ("Сынып / Пән",        "1F3864"),
            ("Толтырылмаған күндер","C0392B"),
            ("Күй",                "1F3864"),
        ]):
            _cell(hdr, i, txt, bg=bg, bold=True, center=True)
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

        for res in problems:
            if res["error"]:
                row = tbl.add_row().cells
                _cell(row, 0, res["name"],    center=False)
                _cell(row, 1, "—")
                _cell(row, 2, res["error"],   bg="FFE0E0", center=False)
                _cell(row, 3, "⚠️ Қате",      bg="FFE0E0")
            else:
                first = True
                for entry in res["unfilled"]:
                    row = tbl.add_row().cells
                    _cell(row, 0, res["name"] if first else "",
                          center=False, bold=first)
                    _cell(row, 1, entry["subject"], center=False)
                    dates_str = ", ".join(entry["dates"])
                    _cell(row, 2, dates_str, bg="FFE0E0", center=False)
                    _cell(row, 3, "❌ Толтырылмаған", bg="FFE0E0")
                    first = False

        # Ширины колонок
        for row in tbl.rows:
            row.cells[0].width = Cm(5.0)
            row.cells[1].width = Cm(4.5)
            row.cells[2].width = Cm(5.5)
            row.cells[3].width = Cm(3.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Полный список всех учителей ────────────────────────────────────────────
    doc.add_page_break()
    t2 = doc.add_heading("", level=2)
    t2.add_run("Барлық мұғалімдер тізімі").bold = True
    t2.paragraph_format.space_after = Pt(4)

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    for i, (txt, bg) in enumerate([
        ("Мұғалімнің аты-жөні","1F3864"),
        ("Логин (ЖСН)",         "1F3864"),
        ("Нәтиже",              "1F3864"),
    ]):
        _cell(hdr2, i, txt, bg=bg, bold=True)
        hdr2[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    for res in results:
        row = tbl2.add_row().cells
        _cell(row, 0, res["name"], center=False)
        _cell(row, 1, res["login"])
        if res["error"]:
            _cell(row, 2, "⚠️ Қате: " + res["error"], bg="FFF3CD", center=False)
        elif res["unfilled"]:
            cnt = sum(len(e["dates"]) for e in res["unfilled"])
            _cell(row, 2, f"❌ {cnt} күн толтырылмаған", bg="FFE0E0")
        else:
            _cell(row, 2, "✅ Толтырылған", bg="E2EFDA")

    for row in tbl2.rows:
        row.cells[0].width = Cm(7.0)
        row.cells[1].width = Cm(4.0)
        row.cells[2].width = Cm(7.0)

    # ── Сохраняем в байты ──────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Главная функция ───────────────────────────────────────────────────────────
def main():
    now_aktau = datetime.utcnow() + timedelta(hours=5)
    print("=" * 50)
    print(f"  Тексеру: {now_aktau.strftime('%d.%m.%Y %H:%M')} (Ақтау)")
    print("=" * 50)

    try:
        teachers = json.loads(TEACHERS_JSON)
    except json.JSONDecodeError as e:
        tg_send(f"❌ TEACHERS_JSON қате: {e}")
        sys.exit(1)

    if not teachers:
        tg_send("⚠️ Мұғалімдер тізімі бос.")
        sys.exit(0)

    # Сообщение о начале
    tg_send(
        f"⏳ <b>Тексеру басталды</b>\n"
        f"👥 {len(teachers)} мұғалім\n"
        f"🕘 {now_aktau.strftime('%d.%m.%Y %H:%M')} (Ақтау)"
    )

    # Проверяем всех учителей
    results = []
    for i, teacher in enumerate(teachers, 1):
        print(f"\n[{i}/{len(teachers)}]")
        results.append(check_teacher(teacher))
        time.sleep(1)

    # Статистика для краткого сообщения
    total   = len(results)
    ok_cnt  = sum(1 for r in results if r["ok"] and not r["error"])
    bad_cnt = sum(1 for r in results if r["unfilled"])
    err_cnt = sum(1 for r in results if r["error"])

    # Краткое сообщение
    lines = [
        f"📋 <b>Журнал тексеру нәтижесі</b>",
        f"🕘 {now_aktau.strftime('%d.%m.%Y %H:%M')} (Ақтау)",
        "",
        f"👥 Тексерілді: <b>{total}</b>",
        f"✅ Толтырылған: <b>{ok_cnt}</b>",
        f"❌ Толтырылмаған: <b>{bad_cnt}</b>",
    ]
    if err_cnt:
        lines.append(f"⚠️ Қате: <b>{err_cnt}</b>")

    bad_list = [r for r in results if r["unfilled"]]
    if bad_list:
        lines.append("")
        lines.append("<b>❌ Толтырылмағандар:</b>")
        for r in bad_list[:10]:   # максимум 10 в сообщении, остальное в Word
            lines.append(f"• <b>{r['name']}</b>")
            for e in r["unfilled"]:
                lines.append(f"  — {e['subject']}: {', '.join(e['dates'])}")
        if len(bad_list) > 10:
            lines.append(f"  <i>...және тағы {len(bad_list)-10} мұғалім (толық есеп Word файлда)</i>")
    else:
        lines.append("")
        lines.append("🎉 <b>Барлығы толтырылған!</b>")

    tg_send("\n".join(lines))

    # Word файл
    print("\nWord есебін жасауда...")
    word_bytes = build_word_report(results, now_aktau)
    filename   = f"journal_report_{now_aktau.strftime('%d%m%Y')}.docx"

    tg_send_document(
        word_bytes, filename,
        caption=f"📄 Толық есеп — {now_aktau.strftime('%d.%m.%Y')}"
    )

    print("\n✓ Аяқталды")


if __name__ == "__main__":
    main()
