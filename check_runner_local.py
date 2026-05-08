"""
check_runner_local.py — Версия для self-hosted runner
Запускается на твоём компьютере через GitHub Actions
Имеет доступ к Chrome/Selenium → делает скриншоты
"""
import io
import json
import os
import sys
import time
import base64
from datetime import datetime, timedelta
from pathlib import Path

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from PIL import Image
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException

# ── Конфигурация из GitHub Secrets / переменных окружения ────────────────────
TELEGRAM_TOKEN   = os.environ.get("TELEGRAM_TOKEN", "")
TELEGRAM_CHAT_ID = os.environ.get("TELEGRAM_CHAT_ID", "")
TEACHERS_JSON    = os.environ.get("TEACHERS_JSON", "[]")

BASE_URL     = "https://mektep.edu.kz"
SCREENSHOT_DIR = Path("screenshots")
SCREENSHOT_DIR.mkdir(exist_ok=True)


# ── Telegram ──────────────────────────────────────────────────────────────────
def tg_send(text: str):
    if not TELEGRAM_TOKEN or not TELEGRAM_CHAT_ID:
        print("[TG]", text[:200])
        return
    url    = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendMessage"
    chunks = [text[i:i+4000] for i in range(0, len(text), 4000)]
    for chunk in chunks:
        try:
            requests.post(url, json={
                "chat_id":    TELEGRAM_CHAT_ID,
                "text":       chunk,
                "parse_mode": "HTML",
            }, timeout=15)
        except Exception as e:
            print(f"[TG] {e}")
        time.sleep(0.3)


def tg_send_document(file_bytes: bytes, filename: str, caption: str = ""):
    if not TELEGRAM_TOKEN or not TELEGRAM_CHAT_ID:
        Path(filename).write_bytes(file_bytes)
        print(f"[TG] Сохранён локально: {filename}")
        return
    url = f"https://api.telegram.org/bot{TELEGRAM_TOKEN}/sendDocument"
    try:
        resp = requests.post(url, data={
            "chat_id":    TELEGRAM_CHAT_ID,
            "caption":    caption,
            "parse_mode": "HTML",
        }, files={
            "document": (filename, file_bytes,
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document")
        }, timeout=60)
        if resp.ok:
            print(f"[TG] Жіберілді: {filename}")
        else:
            print(f"[TG] Қате: {resp.text[:200]}")
    except Exception as e:
        print(f"[TG] Exception: {e}")


# ── Chrome / Selenium ─────────────────────────────────────────────────────────
def make_driver() -> webdriver.Chrome:
    opts = Options()
    # self-hosted runner — компьютер есть, но окно не нужно показывать
    opts.add_argument("--headless=new")
    opts.add_argument("--no-sandbox")
    opts.add_argument("--disable-dev-shm-usage")
    opts.add_argument("--window-size=1600,900")
    opts.add_argument("--disable-gpu")
    opts.add_argument("--lang=ru-RU")
    opts.add_experimental_option("excludeSwitches", ["enable-automation"])
    opts.add_experimental_option("useAutomationExtension", False)
    # ChromeDriver должен быть в PATH (установлен на компе)
    try:
        driver = webdriver.Chrome(options=opts)
    except WebDriverException:
        # Пробуем явный путь для Windows
        service = Service("C:/chromedriver/chromedriver.exe")
        driver  = webdriver.Chrome(service=service, options=opts)
    driver.set_page_load_timeout(30)
    return driver


def selenium_login(driver, login_val: str, password: str) -> bool:
    try:
        driver.get(BASE_URL + "/ru/login")
        wait = WebDriverWait(driver, 10)
        wait.until(EC.presence_of_element_located((By.NAME, "login")))
        driver.find_element(By.NAME, "login").send_keys(login_val)
        driver.find_element(By.NAME, "password").send_keys(password)
        driver.find_element(By.CSS_SELECTOR, "button[type=submit]").click()
        time.sleep(2)
        return "/login" not in driver.current_url
    except Exception as e:
        print(f"    [login] {e}")
        return False


def get_subjects_selenium(driver) -> list:
    try:
        driver.get(BASE_URL + "/ru/grades")
        time.sleep(1)
        soup     = BeautifulSoup(driver.page_source, "html.parser")
        subjects = []
        for a in soup.select("a[href*='/ru/grades/']"):
            href = a.get("href", "")
            txt  = " ".join(a.get_text().split())
            if href not in ("/ru/grades", "") and txt:
                subjects.append({
                    "name": txt,
                    "url":  BASE_URL + href if href.startswith("/") else href
                })
        return subjects
    except Exception as e:
        print(f"    [subjects] {e}")
        return []


def take_screenshot(driver, url: str, name: str) -> str | None:
    """Открывает страницу и делает скриншот. Возвращает путь к файлу."""
    try:
        driver.get(url)
        time.sleep(2)   # ждём загрузки

        # Скроллим вниз чтобы видеть всю таблицу
        driver.execute_script("window.scrollTo(0, 300);")
        time.sleep(0.5)

        # Делаем полный скриншот страницы
        safe_name = "".join(c if c.isalnum() or c in "-_" else "_" for c in name)
        path      = SCREENSHOT_DIR / f"{safe_name}.png"
        driver.save_screenshot(str(path))

        # Сжимаем для уменьшения размера Word файла
        img = Image.open(str(path))
        img = img.convert("RGB")
        # Уменьшаем до ширины 1200px если больше
        if img.width > 1200:
            ratio  = 1200 / img.width
            new_h  = int(img.height * ratio)
            img    = img.resize((1200, new_h), Image.LANCZOS)
        img.save(str(path), "PNG", optimize=True)

        print(f"    📸 Скриншот: {path.name}")
        return str(path)
    except Exception as e:
        print(f"    [screenshot] {e}")
        return None


def check_journal_filled(driver, url: str, days_back: int = 14) -> list:
    """Проверяет незаполненные даты через Selenium."""
    today    = datetime.now()
    unfilled = []
    try:
        driver.get(url)
        time.sleep(1.5)
        soup  = BeautifulSoup(driver.page_source, "html.parser")
        table = soup.find("table")
        if not table:
            return []
        headers   = table.find_all("th")
        date_cols = {}
        for idx, th in enumerate(headers):
            txt = th.get_text(strip=True)
            for fmt in ["%d.%m", "%d.%m.%Y"]:
                try:
                    dt = datetime.strptime(
                        txt + (f".{today.year}" if len(txt) <= 5 else ""), "%d.%m.%Y")
                    if 0 <= (today - dt).days <= days_back:
                        date_cols[idx] = txt
                    break
                except ValueError:
                    continue
        rows = table.find_all("tr")[1:]
        for col_idx, date_str in date_cols.items():
            col_cells = []
            for row in rows:
                cells = row.find_all(["td","th"])
                if col_idx < len(cells):
                    col_cells.append(cells[col_idx].get_text(strip=True))
            if col_cells and all(c in ("","-","н") for c in col_cells):
                unfilled.append(date_str)
    except Exception as e:
        print(f"    [check_journal] {e}")
    return unfilled


def check_teacher(teacher: dict) -> dict:
    name     = teacher.get("name", "—")
    login_v  = str(teacher.get("login", ""))
    password = str(teacher.get("password", ""))
    print(f"\n  👤 {name} ({login_v})")

    result = {
        "name":        name,
        "login":       login_v,
        "ok":          False,
        "error":       None,
        "unfilled":    [],
        # unfilled item: {"subject": str, "dates": [str], "screenshot": str|None}
    }

    driver = make_driver()
    try:
        if not selenium_login(driver, login_v, password):
            result["error"] = "Кіру мүмкін болмады"
            print("    ❌ Кіру қатесі")
            return result

        print("    ✓ Кірді")
        subjects = get_subjects_selenium(driver)
        if not subjects:
            result["error"] = "Пәндер табылмады"
            return result

        print(f"    Пәндер: {len(subjects)}")
        all_ok = True

        for subj in subjects:
            dates = check_journal_filled(driver, subj["url"])
            if dates:
                all_ok = False
                # Делаем скриншот незаполненного журнала
                shot_name = f"{login_v}_{subj['name'][:20]}"
                shot_path = take_screenshot(driver, subj["url"], shot_name)
                result["unfilled"].append({
                    "subject":    subj["name"],
                    "dates":      dates,
                    "screenshot": shot_path,
                })
                print(f"    ✘ {subj['name']}: {', '.join(dates)}")
            else:
                print(f"    ✓ {subj['name']}")
            time.sleep(0.5)

        result["ok"] = all_ok

    finally:
        try:
            driver.quit()
        except Exception:
            pass

    return result


# ── Word отчёт ────────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _cell(cells, i, text, bg=None, bold=False, center=True, size=8, fg=None):
    cells[i].text = ""
    p   = cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold      = bold
    if bg:  _set_cell_bg(cells[i], bg)
    if fg:  run.font.color.rgb = fg


def build_word_report(results: list, check_time: datetime) -> bytes:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Cm(1.5)
    sec.left_margin = sec.right_margin = Cm(1.8)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(9)

    # ── Заголовок ──────────────────────────────────────────────────────────────
    t = doc.add_heading("", level=1)
    r = t.add_run("Электрондық журнал толтыру туралы есеп")
    r.font.size = Pt(13); r.bold = True
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph(
        f"Күні: {check_time.strftime('%d.%m.%Y')}  |  "
        f"Уақыт: {check_time.strftime('%H:%M')} (Ақтау)")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)

    # ── Статистика ─────────────────────────────────────────────────────────────
    total   = len(results)
    ok_cnt  = sum(1 for r in results if r["ok"] and not r["error"])
    bad_cnt = sum(1 for r in results if r["unfilled"])
    err_cnt = sum(1 for r in results if r["error"])

    stbl = doc.add_table(rows=1, cols=4)
    stbl.style = "Table Grid"
    sc = stbl.rows[0].cells
    for i, (lbl, val, bg) in enumerate([
        ("Барлығы",           str(total),   "4472C4"),
        ("✅ Толтырылған",    str(ok_cnt),  "1B7A3E"),
        ("❌ Толтырылмаған",  str(bad_cnt), "C0392B"),
        ("⚠️ Қате",           str(err_cnt), "E67E22"),
    ]):
        _set_cell_bg(sc[i], bg)
        p2  = sc[i].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        r1  = p2.add_run(lbl + "\n")
        r1.font.size = Pt(8)
        r1.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r2  = p2.add_run(val)
        r2.font.size = Pt(16); r2.bold = True
        r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ── Таблица проблемных ─────────────────────────────────────────────────────
    problems = [r for r in results if r["unfilled"] or r["error"]]
    if problems:
        h = doc.add_paragraph()
        h.add_run(f"Толтырылмаған журналдар — {len(problems)} мұғалім:").bold = True
        h.paragraph_format.space_after = Pt(3)

        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        for i, (txt, bg) in enumerate([
            ("Мұғалімнің аты-жөні", "1F3864"),
            ("Сынып / Пән",         "1F3864"),
            ("Толтырылмаған күндер", "C0392B"),
            ("Күй",                  "1F3864"),
        ]):
            _cell(hdr, i, txt, bg=bg, bold=True,
                  fg=RGBColor(0xFF,0xFF,0xFF))

        for res in problems:
            if res["error"]:
                row = tbl.add_row().cells
                _cell(row, 0, res["name"], center=False)
                _cell(row, 1, "—")
                _cell(row, 2, res["error"], bg="FFE0E0", center=False)
                _cell(row, 3, "⚠️ Қате",   bg="FFE0E0")
            else:
                for idx, entry in enumerate(res["unfilled"]):
                    row = tbl.add_row().cells
                    _cell(row, 0, res["name"] if idx == 0 else "",
                          center=False, bold=(idx==0))
                    _cell(row, 1, entry["subject"], center=False)
                    _cell(row, 2, ", ".join(entry["dates"]),
                          bg="FFE0E0", center=False)
                    _cell(row, 3, "❌", bg="FFE0E0")

        for row in tbl.rows:
            row.cells[0].width = Cm(5.0)
            row.cells[1].width = Cm(4.5)
            row.cells[2].width = Cm(5.5)
            row.cells[3].width = Cm(2.0)

    else:
        ok_p = doc.add_paragraph()
        ok_p.add_run("🎉 Барлық мұғалімдер журналды дұрыс толтырған!").bold = True
        ok_p.runs[0].font.color.rgb = RGBColor(0x1B,0x7A,0x3E)
        ok_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Скриншоты ──────────────────────────────────────────────────────────────
    has_shots = any(
        e.get("screenshot") and Path(e["screenshot"]).exists()
        for res in results
        for e in res.get("unfilled", [])
    )

    if has_shots:
        doc.add_page_break()
        sh = doc.add_heading("", level=1)
        sh.add_run("Скриншоттар — толтырылмаған журналдар").bold = True
        sh.paragraph_format.space_after = Pt(6)

        for res in results:
            if not res.get("unfilled"):
                continue
            teacher_shots = [
                e for e in res["unfilled"]
                if e.get("screenshot") and Path(e["screenshot"]).exists()
            ]
            if not teacher_shots:
                continue

            # Имя учителя
            name_p = doc.add_paragraph()
            name_p.add_run(f"👤 {res['name']}").bold = True
            name_p.runs[0].font.size = Pt(10)
            name_p.paragraph_format.space_after  = Pt(3)
            name_p.paragraph_format.space_before = Pt(8)

            for entry in teacher_shots:
                # Подпись скриншота
                cap = doc.add_paragraph()
                cap.paragraph_format.space_after  = Pt(1)
                cap.paragraph_format.space_before = Pt(0)
                r_cap = cap.add_run(
                    f"📚 {entry['subject']} — "
                    f"толтырылмаған: {', '.join(entry['dates'])}"
                )
                r_cap.italic = True
                r_cap.font.size = Pt(8)
                r_cap.font.color.rgb = RGBColor(0x55,0x55,0x55)

                # Скриншот
                try:
                    doc.add_picture(entry["screenshot"], width=Cm(16.5))
                    doc.paragraphs[-1].paragraph_format.space_after = Pt(6)
                except Exception as ex:
                    doc.add_paragraph(f"[Скриншот қатесі: {ex}]")

    # ── Полный список ──────────────────────────────────────────────────────────
    doc.add_page_break()
    t2 = doc.add_heading("", level=2)
    t2.add_run("Барлық мұғалімдер тізімі").bold = True
    t2.paragraph_format.space_after = Pt(4)

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0].cells
    for i, (txt, bg) in enumerate([
        ("Мұғалімнің аты-жөні", "1F3864"),
        ("Логин (ЖСН)",          "1F3864"),
        ("Нәтиже",               "1F3864"),
    ]):
        _cell(hdr2, i, txt, bg=bg, bold=True,
              fg=RGBColor(0xFF,0xFF,0xFF))

    for res in results:
        row = tbl2.add_row().cells
        _cell(row, 0, res["name"], center=False)
        _cell(row, 1, res["login"])
        if res["error"]:
            _cell(row, 2, "⚠️ " + res["error"], bg="FFF3CD", center=False)
        elif res["unfilled"]:
            cnt = sum(len(e["dates"]) for e in res["unfilled"])
            _cell(row, 2, f"❌ {cnt} күн толтырылмаған", bg="FFE0E0")
        else:
            _cell(row, 2, "✅ Толтырылған", bg="E2EFDA")

    for row in tbl2.rows:
        row.cells[0].width = Cm(7.0)
        row.cells[1].width = Cm(4.0)
        row.cells[2].width = Cm(7.0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Главная функция ───────────────────────────────────────────────────────────
def main():
    now_aktau = datetime.utcnow() + timedelta(hours=5)
    print("=" * 55)
    print(f"  Тексеру: {now_aktau.strftime('%d.%m.%Y %H:%M')} (Ақтау)")
    print("=" * 55)

    try:
        teachers = json.loads(TEACHERS_JSON)
    except json.JSONDecodeError as e:
        tg_send(f"❌ TEACHERS_JSON қате: {e}")
        sys.exit(1)

    if not teachers:
        tg_send("⚠️ Мұғалімдер тізімі бос.")
        sys.exit(0)

    tg_send(
        f"⏳ <b>Тексеру басталды</b>\n"
        f"👥 {len(teachers)} мұғалім\n"
        f"🕘 {now_aktau.strftime('%d.%m.%Y %H:%M')} (Ақтау)"
    )

    results = []
    for i, teacher in enumerate(teachers, 1):
        print(f"\n[{i}/{len(teachers)}]")
        results.append(check_teacher(teacher))
        time.sleep(2)

    # Краткое сообщение
    total   = len(results)
    ok_cnt  = sum(1 for r in results if r["ok"] and not r["error"])
    bad_cnt = sum(1 for r in results if r["unfilled"])
    err_cnt = sum(1 for r in results if r["error"])

    lines = [
        f"📋 <b>Журнал тексеру нәтижесі</b>",
        f"🕘 {now_aktau.strftime('%d.%m.%Y %H:%M')} (Ақтау)",
        "",
        f"👥 Тексерілді: <b>{total}</b>",
        f"✅ Толтырылған: <b>{ok_cnt}</b>",
        f"❌ Толтырылмаған: <b>{bad_cnt}</b>",
    ]
    if err_cnt:
        lines.append(f"⚠️ Қате: <b>{err_cnt}</b>")

    bad_list = [r for r in results if r["unfilled"]]
    if bad_list:
        lines += ["", "<b>❌ Толтырылмағандар:</b>"]
        for r in bad_list[:10]:
            lines.append(f"• <b>{r['name']}</b>")
            for e in r["unfilled"]:
                lines.append(f"  — {e['subject']}: {', '.join(e['dates'])}")
        if len(bad_list) > 10:
            lines.append(f"  <i>...тағы {len(bad_list)-10} (Word файлда)</i>")
    else:
        lines += ["", "🎉 <b>Барлығы толтырылған!</b>"]

    tg_send("\n".join(lines))

    # Word файл со скриншотами
    print("\nWord есебін жасауда...")
    word_bytes = build_word_report(results, now_aktau)
    filename   = f"journal_report_{now_aktau.strftime('%d%m%Y')}.docx"
    tg_send_document(
        word_bytes, filename,
        caption=f"📄 Толық есеп + скриншоттар — {now_aktau.strftime('%d.%m.%Y')}"
    )

    # Чистим скриншоты
    for f in SCREENSHOT_DIR.glob("*.png"):
        try:
            f.unlink()
        except Exception:
            pass

    print("\n✓ Аяқталды")


if __name__ == "__main__":
    main()
